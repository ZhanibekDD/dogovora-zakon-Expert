from __future__ import annotations

import base64
import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, RGBColor

from app.services.master_template_service import build_master_template as build_v3_template

SCHEMA_MARKER = "ZakonExpert contract schema v5 official-brand"
NAVY = "20364F"
GOLD = "B78A43"
PALE_BLUE = "F3F6F9"
PALE_GOLD = "FBF8F1"
WHITE = "FFFFFF"
DARK_FILLS = {"102A43", "173A5E", "171717", "20242A", "000000"}
OFFICIAL_PHONE = "+7 700 309 7566"

# Exact ZakonExpert header logo supplied by the owner. Kept inline so the generated
# template never silently falls back to the temporary shield placeholder.
OFFICIAL_LOGO_PNG_B64 = "iVBORw0KGgoAAAANSUhEUgAAALYAAAAmCAYAAABplmFHAAAbDElEQVR42u2cZ3hVVdbHf/uc2+/NTc9NDwFCCaGDIEVRVEQFFBmxwFhGHSyIBUQUGwPYB1ABBREFUWwUHQVEinQkQMDQSwhpENLb7We/H5JgIAEy7zjO+45Zz3O+3Lv32nuv9d9rr7X2OkdIKflnSEpJWUWlTNt3mO0797Jn32GOZ+dQVl6JSTUSFxtF26REenZtR9eObYl2hAshBEIImqiJfi8SjQG2lBKfX2PNlp/lWzMXcvDgEXpelkz/PpfTPD4au9WIqgJ+DadXI+t0Idt2pbNy3QZUxcKj993BiGEDCbRZxT8FcK1MHtp1iCzneXNU7bTq2oZ4I9XMPKflrtRMirDTulsb4gw1v9f0zxYRdO2cQJCCAI3iY+lyV64bW3w7uidYhAJolbly26Y9pJ+qhAAHHXp047IYc/V/JZlya/ppXOZIunWOJ1BBgJ/cfWnygCuCbp1j0DLS2ZXrop409WF06BpKYdphcmUYHbslElYzj6Kjv8jdpyTRyR1orWSxNf00Z5eqMxDsaEZK8yBhPDvnC/Dv3gz9ifPGF3ps4XGkJIVhUzRRcnyf3JnjrN8fgSWmDb2aB4g/FLD9msaq9Zvl0y9NJykxiof/PISEqCCsRgN6vYJOr8dosqOoRjR8eFzl+F2VeP0St0+SX1TF8h828cnSNdw/8naeGnUnZpOxcQD37JHPDB7H+1naecrsyNQVb/BQlCIASldNkd2eWkcBwdw2bT7vX2sTdfvPFUP4Zvkj9ZoomD7fHn7mM85kjCMj2bez1VhiPytC+RfnlnMpgINg0UPLjceNZSrRj3LBw92xLpzJj3vW8oJJZa/zpnNqz3MAsrlwlG38UT2zXy/5F48b/6FIYvy0M5fQ2B/Fqx5nMg5TzD4g2Kuf3Um824ME5xZLx+8bSrfBg1j+aIH6byvZoxzGOiJ6jmCD/5+K9rMBy7C/ymCpzU0voK9zQCmzXiEqE9GMWhhNv56QlZocfffSX065b8K2LqLWenC0jJ59+jn8XqqmD31cRIiTFisdoIjmglFZ2jQvbDaHWf7IzWCQ/Pkw2F6bh98NQuWbaRL/9v56J1JskeX9pcGt6GteHr+AvmQv3rzVexdxN0TVuPqP5CbIpQaq14iV3y3g5KYjvQmnR++305x//4yWOE85hpFP38s7xzzOUeaDWfhrHu5IkQRWuE6OeGZRWz2d2D8vPE80TNC+HN/lm89PZVps15hYqu5vBUACBW9yGXB9KWMXHinbNeQ5JRIRk57nXHJ6q+WQzURYjZjeXAMozc9xbRp81nR8wEp3p7LsuJYHnh1BD0tCDdIUIi/9QWWj0pC8RaT+vHfefjzT3jxix5MvCh/2HX+/7KKo/+YzQPvrOKlD69i0+NvkfZnH9J/mNfv+xufigF8PG8EHVWB3hb0X+eKKBcC9Z70g7LXwDsZdFUP5kweTbsWDuKSuoqwmNZC1Rsv6TMLIRCKSmBIrIhu3l20bN6M0Xf24b03nuCBsS8y/b1PpaZpl5iegZCoSJEQGyViDUd4d/pqchJvZtYLVxOnVgNXK9zMkm1O4q8awbgrIyjfup7vi+rzLd25gNvHfMbhFney6L1qUING0U8/sqpAoeNdoxnbM0IYAUv0ZWL8+EG0oYDvlm2jDAAjfW7oTdC+r5jyTX4Dlq9anNagMOKiHaL2iXUECgsIzG3F48/fRkrxjzz75CQmfFNAs2GP8Ex3y7leltlOVLRDxCe0ETcN6kKc6iMnq3a8i/A///+YRNH36vbEqhqlZeXo7KHERjtEnMOOVQVUE8EOB3HRESLSbvivC4B0DYF608698v7HX2Tua2NpFmklNrG1UHWmC1p2v89NbsYeGRHTEqMlpJ4lFkJgDggXcbZQDMZD8uuZzzNu6lwKSork38Y/IhRFuYRLckLOmTCdLyva8/K7f+GKQKVmAD/Zq9axyR3Fvdel0FPrSfzilSxZU8Qdw8PO7lpZtJGxTxSSVxHIjbfdwOVBv/bPOpmPSwTQtk30OcIwtGxJGwscyc0lxw9IH+YeIxlzahcTZ3/CmmvubyAmOM2iJ+/mW32tKysIvvYJVj/TXZgAS4fh4uVb1spbPtuLDOzL+492rvHXfyVfeT6Hj56QOm8p27/YRqbfTN928eiOXYx/x2oesoytn81hwmoh0Zyc2L6RvTi4c2BHTPCHit5154M0Lf2g/MuYF/lsxnNEh6hEJnYSQigNAtrjKpO7d6zko0XL+fmIk4QwjTsG95MDrh+KPTSuXj8hFBxxbYSlKFvOeOlhxk5awKQZc+WLjz94YbdEq5Cb3p7K31JN3PLK04xqVce6+HPlspX7cJpbo+1fwSeAw+Jiy4qfyBp2q0yonavTS8Lwu+i9eTFLXnuN2W1ekY+2NghQMOh1COnD7TnPBmtePH5Ap0NfzQVEFHc/PoRP7v6SyR9dzT3nT1lYaXvlAPo56gRmbRxnhawVpckvN55GNZug/BeWrT3D0KEO1DruUtbyV7liea35ttLu5kd45ZYI8t+8NH+km+x9aWww+inOzSLP35JR707i5b5BQvmDZUXOAXZhaZkc/uBY5r4+tgbUnesBTkqNytJ8ueHHJSxcspafs6BMmhDSSl6Wws5ZW5n5xQaG92stbx46ksi41kJRdedY74CQWKGoqnx14ghGjHmDbint5I3X9G4A3H6yV7zNQwtzaDFyKm9eH1EHBOA7vp6lv3jR6TNYPGt+NTR8Cq60n1iSeTNPxNYc0DFXM3HcPXQerCfrgY+YPO592i54RPYPUkWz5ERCxTF2bN5L6YAeMpDqjEXJtp3scilEtm1DvLrt7JimlD/x3E0/MuLTBaxooZ5rBoWNLkPuZEK3Bo52rVSunjGLxafjeHj2w/D6c7z7zhyW9nlODotQzroSsTc9wQd3JaJTDARHRtM8zCTAQz7IC/P31HQP50+TZ/FKN1Wc+Ow52W/qHjbtLMDXN0Ia/qgW269pjBz9PE8+cCfNHKYaSy3OsdDF+Rny62VfsXTFVvYVKLikEUXqMEgFiUQT4MHMkTMmXvkig49WjOeaHs3k3cNvp1W7XkLUuBxCCKyBUSLIWS5nTX6UYaMm0b3zJ9IRFnLusXx0mXxo8jpyg7oyLCaLpV9l1ZhOHc169iPwuw38okVy77sf8kbPamW7tkyXl/11BctWZDL6gXMXa2l/u5jzzBF5/Qvf8ujzLflu+kDZvPfN3JeykVeWT+O+kPsZfWWk9Gdu5d0Za8mztmfyXZ0wFW+r40YHiGsfukv2XzOd1WmgJNbd9S6Ob1/HF3nqr6kmYabV5T1I2P8hE5blk/CnqTzZowvK+JtYMWo5L0/fwpWT+0h7bR4kKIZOHdoIY4N+34X4dz2voUqzWx/ksX88xuRPZvLOgGmMb2v441lsKSUr122RPq+TAX2SiU1s26D78dVHM5i67AhuwkAoKIJq+yaqj0Vx9kAFTZjIqjKxcHU+Jw69yPwFy6TR/GuuVAhBaGQrUVWxS45/+A4ee24qi99745yg1HM4nbQyiUYq0yen1pmJiRunRJDwQyYydghDuvyqNFPXfgyI/p75P6xj793d6ik8fvCTvLP/JHd9OosHZsez7NF2PDltIv5Js3nvw1e5ZW51mi20TR9eePoRHmqpCu+Oc9O/avR1YuLI7+XGdw/V2spaZ54fZr3BD+fEk9H89R2B5/WVnIy4loW1fvVlI+XkIdu4c+n7TL6+A6+bG3M71jD/UfNnc2O9uLu5GDXuZrnkns+ZOXUpg+YNl8mGP47VFlJKvD4fna++Q86a+ijtWkYTGpnUoAC+mvemfGTeHgQ6hJAXvdCpBahfSq5p62Ph+x8KVVffDvk8VfJExn5ueWAKX344nTYtE/5jwneX5MrjeS50IREkOGzij2Xj/gvTfWs275BJzaJpFm4ixNFCXAiskdHhqKoTgQby4qm+ml7ocOMIsSEUtcwNqDXejxUocs4mUOg1YTJZL8k+FRVR0SQpdpY14C5FtTl5eXnWesaK2+QWANCLu3sKzpMJiN6k8DlckuzseZ2TcujYPc/8FXqCDbWFGDrA9Cpbnw+wJ/OsS/uJ9N7BpduIO0TLZTuAVn4JYcz29Ly2mCOrKqs3iPl+ymrSiA2XFcvyq3YNYfCyHtJsczjF4+KPS4ZuWY6O12JUlQcpVwLqxGvhitrFVllKmAmuFtrjG43uqCAs+6ZajBVJ/PRcObvQ5EODEYVZCVerxWLuWZ8JVAYzKr0uZ0N8O2I3VaTSfJW4jfEoK+Nv3VmVFz43C50QYE1NTEqerMd4a3AJy+cFNdyF7J94VLpL89HSRlPuFzBIUM0BrX2vicQvajC55aoaLhz15JTCb5KC6YAJ15XBf4LtBcBNnSKi5LCHPxOMxUGD2q0FeVispf5FO3/B04F0HcgMboL4MeZsYSMMyo4M/DXAXZDeq3HIzZJ6sorKlEViXoJ/xpAKCqhIQEoVNT41TUBZE1OW0gBQkPWqlhIrCYVo8nWmJQ6QlHR6SQ+r6/ObmpDsyHjkd+NIOPkL3iTukDOTspkc5oHKpSoScRdN45IcwBGe7DQyaOylCrOpK4ioMfbhBm+5Uh1xpyKQ+soD7uaYKOgoq5sPdvJ2N+cFsMvQ907D1AwJr8g+kQckBVOPUbjTn75dMvZOy17xwl06WCvWaSP8sh46TqynXJvigxU88j5cRrlibejoMPeaRLxuWPYv2EDYde3JDCkgFNH9+ON7oiucqs8c8qIvVUUZJ3Pt45kwltjrdhIQYGLYIeeqn1vyf0FPUl0JOA+vJESVzsZYizmTMYhlIhrsCoZFzYgEYNpf10/DMZQTFazkAUZ0laxltN55QTFWvFkbKNUn0R0qEI5OgJSxtKpg11oee/JTZ9/Q0GrTtgq1jXQ3k/pjvc4FXQ3PdrvYfumMFoPsHBgySxOtn1Vhl5A9ihJxF8/mThT7cshR+UZ9AR2nUz3DnbhPzZFrl9d27ghvTbAA9AZdSaQGpr0Nwp8EREh6GUBHmGuAbao8bclEhUFJ1Iz41f9KFISZNWjN1oac62GkBK/FAhFnK9ZYvrez5lvXmTDBzaJVyXkiilEmteQiRFDUDzW2kVJQJbgNt1D13aRiIKa4/zYm/LnzQfwcYzNb89D+t1oYhKHg/uBswB9p6nE2BVRfjZ40GMM6yCMgCxNq+N7eShcP4I1m6t3s5Iwmn7XPEbLnL+x44PVGJRKfPYbaX+5jWIAbET0/gvZn8zmcOb7tO1/LyX/eJafDgdKndeJodWTdG5lI2tHfb5X3XCNAFBCh5HcM509Xw4n16KTXnckcQPvJzj6cVrmvcyueT+hNzjxm/vRblA/9O6Mi/iTgVhCYkWt4ylCh5Hcax9py+7glNki/d4gIvtPItJUQTkeCjfcw7qtSM2vYOkwGkd0d4J67a/X3lH1LTv2aMTdOhBz+Z5q3mHDSGrxHb+sGMuxvAvI3r+TAx8O4bBAgp6wK1/kgudNA3qtVu55PK76GDHkz2Plkw8OpkNSFEERLS9567f0s7ny4Zkr0WQAol4uu7oOWviNqMKJX0hubK/y/swPhaq7uA+vaV6OH0yTvYdP4PjP32I1m+vPRTpxl5dJzOHCqP/n6tX8RybJ9Wkp9Bk2VBgFoOXI458/hbP3AtrF/zb1yD7nKen1B2C0WS9RTefCU1YiNVOYMBl0jR/AUyqdLi96W5jQ1RnA7zwtPX5bI8a9CHlLpdNZn/dv0f73kH29/ZvcOpHs08Ukxwc2qkNsVAgKejSpIIUPKSSKVBBSIBD4hIqieJEo+KWRyBAfygVuHc+9WnfJKpcHoSmYjMaGFyvMGO3m/5UgREAKjri4X9+sEBasCb0wWH67gk6dOVI0DqYmDPbIf34dhkBhbuBKVDU7hPlfnbw+UJj1/572v4fs6+nisi7JrNm0jas6x59T43EhCnGEEKxVUIQOTSp13BGQUkFBoqk+NL+CXpYTGRZ1NtC8GJWX5lNc4aNLShLKv+GNdiVyqEiJrCvtYOHo9VjTTcbvcQv4H5C90q1TMj+s34LLq4HULtkhvkUPbr0yHgNVNTUjak0A6cen+lE1wK9DxUeX2EqGDXuwUZ9ecDvL2LgznUE3XNH0qYYm+teBHR0RJoTQkV/soqzk1CXrNPQGi5gwcQpDe0ag4gbhQ5OgaAqqJvELiUH6SHZUMWPKSzji214SpZrmw+3RWLRkFQP7923SShP968AWQvDofXewfNUGygvzGlVdZ7IGi0kvvczNnY3o/T5UCRo+pNCB0GgZWsbMFx4nsVXXRn1mofj0cZl5pgKLyUR8dFSTuW6i3wbYI4YNZNGydRRX+XFVFDaqus5md4ipf3uF69sKVM2HpgC4SbRXMn3i/bTu1K9RoJaan4qyMv4+93NefvpRFKUJ1030GwAbINBmFfeP/BMLv1lLQV4mshG+NoA9JE68/toU+rZwoWqSeGMxM565i449BjcO1FJyJveAPHaqkpOZpxhwVe8mVDfRb0O1ZaaVlU7ZqtdNcv3Gn+SpzHRZW/NxqUfTNE7nHJB3Db9BrluxSGqav1H9pJQ4K4rlkYO7ZadrbpObd6TJxvZrepqeSz2ibj31tt3p8sGxz/PVuy8QGRmGPTiu0Za3vChbBoRECyHURm0ov89F1rFf5N/nr0JnMvHWi0+JpmxIE/2mrghUFy/17Jwi7rvtTzw9dS7FRWVUlubJxgSTQoia74g0HtTZmQfkNxv2k3Y4g1eeHdME6ib6TaneRyk1TeP519+VB/dn8trEPxNoNxPmSLok6spLcmRAULS41Ocr3FWl8lTOMb75KZ1Fy9az+vN3sdusTahuon8vsGvBPWnGHLly1WZmT3mUcIsXoXmwhEYjhUp1mY5ECg0h9QipkZ+VRmh0CopOV11BUvsWuxQo+PG6yqgoK8JvjOTt+d+y50gm3348HbvN1gTqJvp9gF3rN3/34xb52MTXGPfQ7VzbqzUWvUZ4TDI6veWc74PUthfnvQcpAE36Kc4/KstLyziR52bs5Jlc0bcbr0wYI4wGfZMGmuj3BXYtQPMLi+XjE6eSvu8oz465l67tYrEZQDWZsAdFYTDZhKLoqC1flVLD53HKstJ83M4yPG5JVn4Vb877lIzsTOa8NoWeXTo2+dRN9J8Ddl2AHzp+Ur78xmw2b93G9f37MOjavsRG2DEbFEwmQ/V3QyR4PR6qXG6KK71sTk1n4VcrsVqtvPz0I1zXr7fQqUpTLUgT/d8Adl2Al5ZXyA1bU/l6xQZ+2rQZp9OJ3mRAr+oAP84qP2g6OnVow+Ab+nLD1X2Ij44SQtAE6Cb6vwnsugCvDh/B7fHicrmlz+tDKAKT2YjZaBTKee9ENlET/Z70P8/uH6EZl84jAAAAAElFTkSuQmCC"


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    return shd.get(qn("w:fill")) if shd is not None else None


def _set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text_color(cell, color: str) -> None:
    rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = NAVY) -> None:
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        run = paragraph.runs[0]
        run.text = text
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for part in (section.header, section.footer):
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _lighten_tables(doc) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fill = (_cell_fill(cell) or "").upper()
                if fill in DARK_FILLS:
                    _set_cell_fill(cell, PALE_BLUE)
                    _set_cell_text_color(cell, NAVY)


def _replace_text(doc, old: str, new: str) -> None:
    for paragraph in _all_paragraphs(doc):
        if old in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old, new)


def _replace_header_branding(doc) -> None:
    logo_bytes = base64.b64decode(OFFICIAL_LOGO_PNG_B64)
    for section in doc.sections:
        header = section.header
        if not header.tables:
            continue
        table = header.tables[0]
        if len(table.rows[0].cells) < 3:
            continue
        row = table.rows[0]
        combined = row.cells[0].merge(row.cells[1])
        paragraphs = list(combined.paragraphs)
        for paragraph in paragraphs:
            paragraph.clear()
        for paragraph in paragraphs[1:]:
            paragraph._element.getparent().remove(paragraph._element)
        paragraph = combined.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(io.BytesIO(logo_bytes), width=Mm(50))


def _rewrite_payment_table(doc) -> None:
    target = None
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "{{ executor_bank_beneficiary }}" in text or "{{ executor_bank_iban }}" in text:
            target = table
            break
    if target is None:
        return

    rows = [
        (
            "ОСНОВНОЙ СПОСОБ",
            "Оплата по счёту или платёжной ссылке, выставленной ZakonExpert для конкретного договора.",
        ),
        (
            "ИДЕНТИФИКАЦИЯ",
            "В назначении платежа или платёжном документе должны быть указаны номер договора и сумма.",
        ),
        (
            "ПОДТВЕРЖДЕНИЕ",
            "Факт оплаты подтверждается банковским, кассовым либо иным платёжным документом, позволяющим установить получателя, плательщика, сумму и дату.",
        ),
        (
            "КОНТАКТ",
            f"Для получения счёта и подтверждения оплаты: {OFFICIAL_PHONE} · {{{{ executor_website }}}}.",
        ),
    ]
    for row, (label, value) in zip(target.rows, rows, strict=False):
        cells = row.cells
        value_cell = cells[1].merge(cells[3]) if len(cells) >= 4 else cells[-1]
        _set_cell_fill(cells[0], PALE_GOLD)
        _set_cell_text(cells[0], label, bold=True, color=GOLD)
        _set_cell_text(value_cell, value, color=NAVY)
        for extra in cells[2:]:
            if extra is not value_cell:
                _set_cell_fill(extra, WHITE)


def _rewrite_security_notice(doc) -> None:
    for table in doc.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "БЕЗОПАСНОСТЬ ОПЛАТЫ" not in text:
            continue
        for row in table.rows:
            for cell in row.cells:
                if "БЕЗОПАСНОСТЬ ОПЛАТЫ" in cell.text:
                    _set_cell_fill(cell, PALE_GOLD)
                    _set_cell_text(
                        cell,
                        "БЕЗОПАСНОСТЬ ОПЛАТЫ: оплачивайте только по счёту или платёжной ссылке, полученной от ZakonExpert по этому договору. Не переводите деньги по реквизитам, которых нет в счёте или подтверждённой платёжной ссылке. ZakonExpert не запрашивает SMS-коды, пароли и доступ к банковскому приложению.",
                        color=NAVY,
                    )
                    return


def build_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    build_v3_template(output_path)
    doc = Document(str(output_path))
    doc.core_properties.subject = SCHEMA_MARKER
    doc.core_properties.title = "Договор оказания услуг ZakonExpert — official brand"

    _lighten_tables(doc)
    _replace_header_branding(doc)
    _replace_text(doc, "{{ executor_phone }}", OFFICIAL_PHONE)
    _replace_text(doc, "ПЛАТЁЖНЫЕ РЕКВИЗИТЫ · сверяйте перед оплатой", "ПОРЯДОК ОПЛАТЫ · сверяйте перед оплатой")
    _replace_text(doc, "ПЛАТЁЖНЫЕ РЕКВИЗИТЫ", "ПОРЯДОК ОПЛАТЫ")
    _replace_text(
        doc,
        "{{ executor_payment_details }}",
        "по счёту или платёжной ссылке, письменно подтверждённой Исполнителем",
    )
    _rewrite_payment_table(doc)
    _rewrite_security_notice(doc)

    doc.save(str(output_path))
    return output_path


def ensure_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    if output_path.exists():
        try:
            current = Document(str(output_path))
            if current.core_properties.subject == SCHEMA_MARKER:
                return output_path
        except (OSError, ValueError, KeyError):
            pass
    return build_master_template(output_path)
