from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

SCHEMA_MARKER = "ZakonExpert contract schema v3 premium"
FONT = "Times New Roman"
UI_FONT = "Arial"
NAVY = "102A43"
NAVY_2 = "173A5E"
INK = "171717"
GOLD = "B58A3A"
MUTED = "667085"
PALE = "F5F7FA"
PALE_BLUE = "F0F4F8"
PALE_GOLD = "FBF7EE"
BORDER = "D1D8E0"
WHITE = "FFFFFF"
PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_LR_MM = 16
MARGIN_TB_MM = 15
CONTENT_W_MM = PAGE_W_MM - MARGIN_LR_MM * 2


def _font(run, *, size=9.4, bold=False, italic=False, color=INK, family=FONT) -> None:
    run.font.name = family
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(key), family)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, *, top=80, start=110, bottom=80, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def _set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def _table(table, widths_mm: list[float], *, border=BORDER, size=5, inner=True) -> None:
    widths = [round(width / 25.4 * 1440) for width in widths_mm]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        if not inner and edge in ("insideH", "insideV"):
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), border)
            el.set(qn("w:space"), "0")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Twips(width)
            _cell_margins(cell)


def _no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def _keep(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True


def _paragraph_rule(paragraph, *, color=GOLD, size=8, space=2) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    _font(run, size=7.1, color=MUTED, family=UI_FONT)




def _set_defaults(doc) -> None:
    section = doc.sections[0]
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.left_margin = Mm(MARGIN_LR_MM)
    section.right_margin = Mm(MARGIN_LR_MM)
    section.top_margin = Mm(MARGIN_TB_MM)
    section.bottom_margin = Mm(MARGIN_TB_MM)
    section.header_distance = Mm(5.5)
    section.footer_distance = Mm(6.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.8)
    normal.paragraph_format.line_spacing = 1.02


def _header_footer(doc) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    table = header.add_table(rows=1, cols=3, width=Mm(CONTENT_W_MM))
    _table(table, [14, 104, CONTENT_W_MM - 118], border=WHITE, size=0)
    _no_borders(table)
    for cell in table.rows[0].cells:
        _cell_margins(cell, top=0, start=0, bottom=0, end=0)
    mark, brand, ref = table.rows[0].cells
    mark.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Left intentionally blank here: this is only ever an intermediate build stage - the
    # light-brand post-processor (master_template_light_service) always merges this cell
    # with `brand` and embeds the real ZakonExpert logo asset before a document is saved,
    # or the whole build fails loudly if that asset is missing. No placeholder icon needed.
    bp = brand.paragraphs[0]
    bp.paragraph_format.space_after = Pt(0)
    _font(bp.add_run("ZAKONEXPERT\n"), size=9.1, bold=True, color=NAVY, family=UI_FONT)
    _font(bp.add_run("ЮРИДИЧЕСКОЕ СОПРОВОЖДЕНИЕ"), size=6.2, bold=True, color=GOLD, family=UI_FONT)
    rp = ref.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    _font(rp.add_run("ДОГОВОР № {{ contract_number }}\n"), size=7.5, bold=True, color=NAVY, family=UI_FONT)
    _font(rp.add_run("{{ contract_date }}"), size=6.6, color=MUTED, family=UI_FONT)
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(0)
    _paragraph_rule(rule, color=GOLD, size=8, space=0)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    footer_table = footer.add_table(rows=1, cols=3, width=Mm(CONTENT_W_MM))
    _table(footer_table, [92, 58, CONTENT_W_MM - 150], border=WHITE, size=0)
    _no_borders(footer_table)
    for cell in footer_table.rows[0].cells:
        _cell_margins(cell, top=0, start=0, bottom=0, end=0)
    left, center, right = footer_table.rows[0].cells
    _font(left.paragraphs[0].add_run("ТОО «ZakonExpert» · БИН {{ executor_identifier }}"), size=6.6, color=MUTED, family=UI_FONT)
    center.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(center.paragraphs[0].add_run("{{ executor_phone }} · {{ executor_website }}"), size=6.6, color=MUTED, family=UI_FONT)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(right.paragraphs[0].add_run("СТР. "), size=6.6, color=MUTED, family=UI_FONT)
    _field(right.paragraphs[0], "PAGE")
    _font(right.paragraphs[0].add_run(" / "), size=6.6, color=MUTED, family=UI_FONT)
    _field(right.paragraphs[0], "NUMPAGES")


def _part(doc, roman: str, title: str, subtitle: str) -> None:
    banner = doc.add_table(rows=1, cols=2)
    _table(banner, [31, CONTENT_W_MM - 31], border=WHITE, size=0)
    _no_borders(banner)
    left, right = banner.rows[0].cells
    _shade(left, NAVY)
    _shade(right, PALE_GOLD)
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    _cell_margins(left, top=55, start=50, bottom=55, end=50)
    _cell_margins(right, top=55, start=120, bottom=55, end=120)
    _font(left.paragraphs[0].add_run(f"ЧАСТЬ {roman}"), size=7.2, bold=True, color=WHITE, family=UI_FONT)
    _font(right.paragraphs[0].add_run(title + "\n"), size=8.7, bold=True, color=NAVY, family=UI_FONT)
    _font(right.paragraphs[0].add_run(subtitle), size=6.7, color=MUTED, family=UI_FONT)
    _keep(banner.rows[0])


def _section(doc, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    _table(table, [11.5, CONTENT_W_MM - 11.5], border=WHITE, size=0)
    _no_borders(table)
    num, text = table.rows[0].cells
    _shade(num, NAVY)
    num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    num.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _font(num.paragraphs[0].add_run(number.zfill(2)), size=7.1, bold=True, color=WHITE, family=UI_FONT)
    tp = text.paragraphs[0]
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(0)
    _font(tp.add_run(title), size=9.0, bold=True, color=NAVY, family=UI_FONT)
    _paragraph_rule(tp, color=GOLD, size=6, space=1)
    for cell in table.rows[0].cells:
        _cell_margins(cell, top=45, start=80, bottom=45, end=80)
    _keep(table.rows[0])


def _clause(doc, number: str, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Mm(0)
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.05)
    p.paragraph_format.line_spacing = 1.03
    _font(p.add_run(number + " "), size=9.15, bold=True, color=NAVY)
    _font(p.add_run(text), size=9.15)


def _title_page(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    _font(p.add_run("ИНДИВИДУАЛЬНЫЙ ДОГОВОР"), size=7.3, bold=True, color=GOLD, family=UI_FONT)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(1)
    title.paragraph_format.space_after = Pt(1)
    _font(title.add_run("ДОГОВОР ОКАЗАНИЯ УСЛУГ"), size=16.3, bold=True, color=NAVY, family=UI_FONT)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    _font(sub.add_run("Консультационное и документальное сопровождение"), size=8.0, color=MUTED, family=UI_FONT)

    meta = doc.add_table(rows=1, cols=3)
    _table(meta, [CONTENT_W_MM / 3] * 3, border=BORDER, size=5)
    labels = (
        ("НОМЕР ДОГОВОРА", "№ {{ contract_number }}"),
        ("МЕСТО ЗАКЛЮЧЕНИЯ", "{{ contract_city }}"),
        ("ДАТА", "{{ contract_date }} г."),
    )
    for cell, (label, value) in zip(meta.rows[0].cells, labels, strict=True):
        _shade(cell, PALE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(cell.paragraphs[0].add_run(label + "\n"), size=6.2, bold=True, color=GOLD, family=UI_FONT)
        _font(cell.paragraphs[0].add_run(value), size=8.6, bold=True, color=NAVY, family=UI_FONT)
        _cell_margins(cell, top=70, start=80, bottom=70, end=80)
    _keep(meta.rows[0])


def _intro_card(doc) -> None:
    card = doc.add_table(rows=1, cols=2)
    _table(card, [CONTENT_W_MM / 2, CONTENT_W_MM / 2], border=BORDER, size=5)
    left, right = card.rows[0].cells
    _shade(left, WHITE)
    _shade(right, WHITE)
    lp, rp = left.paragraphs[0], right.paragraphs[0]
    _font(lp.add_run("ИСПОЛНИТЕЛЬ\n"), size=6.4, bold=True, color=GOLD, family=UI_FONT)
    _font(lp.add_run("{{ executor_brand_name }}\n"), size=8.7, bold=True, color=NAVY, family=UI_FONT)
    _font(lp.add_run("{{ executor_identifier_label }} {{ executor_identifier }} · {{ executor_director_name }}"), size=7.7, color=MUTED)
    _font(rp.add_run("КЛИЕНТ\n"), size=6.4, bold=True, color=GOLD, family=UI_FONT)
    _font(rp.add_run("{{ client_full_name }}\n"), size=8.7, bold=True, color=NAVY, family=UI_FONT)
    _font(rp.add_run("ИИН {{ client_iin }} · {{ client_phone }}"), size=7.7, color=MUTED)
    for cell in (left, right):
        _cell_margins(cell, top=75, start=120, bottom=75, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _keep(card.rows[0])
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.space_before = Pt(2)
    intro.paragraph_format.space_after = Pt(3)
    _font(
        intro.add_run(
            "{{ executor_full_name }}, {{ executor_identifier_label }} {{ executor_identifier }}, в лице руководителя "
            "{{ executor_director_name }}, действующего на основании Устава, именуемое «Исполнитель», и "
            "{{ client_full_name }}, ИИН {{ client_iin }}, именуемый(ая) «Клиент», совместно — «Стороны», "
            "заключили настоящий договор на следующих условиях."
        ),
        size=8.75,
    )


def _key_terms(doc) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(1.5)
    _font(title.add_run("КЛЮЧЕВЫЕ УСЛОВИЯ"), size=8.3, bold=True, color=NAVY, family=UI_FONT)
    _font(title.add_run("  ·  что получает Клиент"), size=7.2, color=MUTED, family=UI_FONT)
    _paragraph_rule(title, color=GOLD, size=8, space=2)

    table = doc.add_table(rows=5, cols=3)
    _table(table, [11, 38, CONTENT_W_MM - 49], border=BORDER, size=5)
    rows = [
        ("01", "УСЛУГА", "{{ service_subject }}"),
        ("02", "РЕЗУЛЬТАТ", "{{ result_definition }}"),
        ("03", "СРОК", "{{ work_period }}"),
        ("04", "СТОИМОСТЬ", "{{ amount_digits }} ({{ amount_words }})"),
        ("05", "ОПЛАТА", "{{ payment_terms }}"),
    ]
    for idx, (row, values) in enumerate(zip(table.rows, rows, strict=True)):
        num, label, value = row.cells
        _shade(num, NAVY if idx in (0, 3) else PALE_BLUE)
        _shade(label, PALE_GOLD)
        _shade(value, WHITE)
        num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        num.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _font(num.paragraphs[0].add_run(values[0]), size=6.6, bold=True, color=WHITE if idx in (0, 3) else NAVY, family=UI_FONT)
        _font(label.paragraphs[0].add_run(values[1]), size=6.8, bold=True, color=GOLD, family=UI_FONT)
        _font(value.paragraphs[0].add_run(values[2]), size=8.45, bold=idx == 3, color=NAVY if idx == 3 else INK)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell, top=62, start=95, bottom=62, end=95)
        _set_row_height(row, 340 if idx != 1 else 410)
        _keep(row)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    note.paragraph_format.space_before = Pt(1)
    note.paragraph_format.space_after = Pt(2)
    _font(note.add_run("Ключевые условия являются неотъемлемой частью договора."), size=6.8, italic=True, color=MUTED)


def _page1(doc) -> None:
    _title_page(doc)
    _intro_card(doc)
    _part(doc, "I", "ИНДИВИДУАЛЬНЫЕ УСЛОВИЯ", "Объём работы, результат, срок и стоимость")
    _key_terms(doc)
    _section(doc, "1", "ПРЕДМЕТ И РЕЗУЛЬТАТ")
    _clause(doc, "1.1.", "{{ service_subject }}")
    _clause(doc, "1.2.", "В состав услуг входят: {{ service_actions }}.")
    _clause(doc, "1.3.", "Проверяемый результат оказания услуг: {{ result_definition }}. Исполнитель подтверждает выполненные действия документами, талонами, ответами органов, перепиской либо иными объективными материалами.")
    _section(doc, "2", "ПОРЯДОК И СРОКИ")
    _clause(doc, "2.1.", "Исполнитель начинает работу после получения необходимых документов и сведений Клиента, а при предоплате — после её поступления.")
    _clause(doc, "2.2.", "Срок действий Исполнителя: {{ work_period }}. Время рассмотрения обращений судом, ЧСИ, нотариусом, банком, взыскателем и государственными органами не включается в этот срок; Исполнитель информирует Клиента о таких периодах ожидания.")
    _clause(doc, "2.3.", "Клиент своевременно передаёт уведомления, постановления, судебные извещения и другие материалы по делу. Задержка предоставления необходимых данных соразмерно продлевает срок работы.")
    _section(doc, "3", "СТОИМОСТЬ И ОПЛАТА")
    _clause(doc, "3.1.", "Стоимость услуг составляет {{ amount_digits }} ({{ amount_words }}). Увеличение согласованной стоимости без письменного согласия Клиента не допускается.")
    _clause(doc, "3.2.", "{{ payment_terms }}")
    _clause(doc, "3.3.", "Оплата производится {{ executor_payment_details }}. Факт оплаты подтверждается банковским чеком, квитанцией либо иным платёжным документом.")
    _clause(doc, "3.4.", "{{ penalty_clause }}")


def _page2(doc) -> None:
    doc.add_page_break()
    _part(doc, "II", "ПРАВИЛА РАБОТЫ И ГАРАНТИИ", "Права, обязанности, приёмка и электронное взаимодействие")
    _section(doc, "4", "ОБЯЗАННОСТИ ИСПОЛНИТЕЛЯ")
    _clause(doc, "4.1.", "Исполнитель обязуется добросовестно анализировать материалы, готовить согласованные документы, соблюдать применимые сроки своих действий и информировать Клиента о существенных этапах.")
    _clause(doc, "4.2.", "По запросу Клиента Исполнитель предоставляет копии подготовленных и направленных документов, полученных ответов и подтверждений подачи.")
    _clause(doc, "4.3.", "Исполнитель выбирает законные способы сопровождения в пределах предмета договора и вправе запросить дополнительные сведения, необходимые для качественного выполнения поручения.")
    _clause(doc, "4.4.", "Решение суда, ЧСИ, нотариуса, банка, взыскателя или государственного органа принимается соответствующим лицом самостоятельно. Исполнитель отвечает за качество и своевременность собственных действий в согласованном объёме.")
    _section(doc, "5", "ОБЯЗАННОСТИ КЛИЕНТА И ПРИЁМКА")
    _clause(doc, "5.1.", "Клиент предоставляет полные и достоверные сведения, сообщает о платежах, соглашениях, судебных актах и иных обстоятельствах, влияющих на дело, и своевременно подписывает документы, когда требуется его личная подпись или ЭЦП.")
    _clause(doc, "5.2.", "Исполнитель направляет Клиенту результат и подтверждающие материалы. Клиент вправе в течение 5 рабочих дней направить конкретные мотивированные замечания. При их отсутствии соответствующий объём услуг считается принятым.")
    _clause(doc, "5.3.", "Подтверждённые недостатки подготовленных Исполнителем документов устраняются без дополнительной оплаты в разумный срок.")
    _section(doc, "6", "ПРЕКРАЩЕНИЕ, РАСХОДЫ И ОТВЕТСТВЕННОСТЬ")
    _clause(doc, "6.1.", "Клиент вправе отказаться от договора. Стороны производят расчёт исходя из фактически выполненных согласованных действий и документально подтверждённых расходов; неиспользованный остаток оплаты возвращается после сверки.")
    _clause(doc, "6.2.", "Обязательные платежи третьим лицам — государственная пошлина, нотариальные, почтовые и иные согласованные расходы — оплачиваются отдельно только после предварительного уведомления Клиента, если иное прямо не включено в стоимость.")
    _clause(doc, "6.3.", "Исполнитель вправе приостановить работу, если без документов или действий Клиента продолжение объективно невозможно, предварительно сообщив, что требуется и как это влияет на срок.")
    _clause(doc, "6.4.", "Исполнитель вправе прекратить договор при требовании незаконных действий, заведомо недостоверных сведениях или существенном повторном нарушении обязанностей Клиентом после письменного предупреждения и предоставления разумного срока для устранения нарушения.")
    _section(doc, "7", "ПЕРСОНАЛЬНЫЕ ДАННЫЕ И ЭЛЕКТРОННОЕ ВЗАИМОДЕЙСТВИЕ")
    _clause(doc, "7.1.", "Клиент даёт согласие на сбор, обработку и хранение персональных данных в объёме, необходимом для исполнения договора, подготовки и направления документов, связи и хранения доказательств выполненной работы.")
    _clause(doc, "7.2.", "Исполнитель принимает разумные меры защиты персональных данных и не передаёт их третьим лицам, кроме случаев, необходимых для исполнения поручения или предусмотренных законом.")
    _clause(doc, "7.3.", "Переписка, скан-копии и сообщения, позволяющие определить отправителя, содержание и дату, могут использоваться как доказательство согласований и уведомлений. Документ с ЭЦП применяется в случаях, установленных законодательством Республики Казахстан.")
    _section(doc, "8", "ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
    _clause(doc, "8.1.", "Договор вступает в силу с даты подписания на бумаге, ЭЦП либо простой электронной подписью через персональную страницу ZakonExpert с явным подтверждением согласия и действует до полного исполнения обязательств.")
    _clause(doc, "8.2.", "Изменения действительны при письменном согласовании, включая электронную переписку, позволяющую определить Стороны и содержание изменения.")
    _clause(doc, "8.3.", "Недействительность одного положения не влечёт недействительность остальных положений договора. Каждая Сторона вправе сохранить и получить свою копию договора.")


def _card_lines(cell, title: str, items: list[tuple[str, bool]]) -> None:
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(title + "\n"), size=6.5, bold=True, color=GOLD, family=UI_FONT)
    for idx, (text, bold) in enumerate(items):
        if idx:
            p.add_run("\n")
        _font(p.add_run(text), size=8.0, bold=bold, color=NAVY if bold else INK)


def _payment_cell(cell, label: str, value: str, *, emphasize=False) -> None:
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _font(p.add_run(label.upper() + "\n"), size=6.2, bold=True, color=GOLD, family=UI_FONT)
    _font(p.add_run(value), size=8.35, bold=emphasize, color=NAVY if emphasize else INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell, top=75, start=115, bottom=75, end=115)


def _page3(doc) -> None:
    doc.add_page_break()
    hero = doc.add_table(rows=1, cols=3)
    _table(hero, [CONTENT_W_MM / 3] * 3, border=BORDER, size=5)
    left, middle, right = hero.rows[0].cells
    _shade(left, NAVY)
    _shade(middle, PALE_GOLD)
    _shade(right, PALE)
    for cell in hero.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _cell_margins(cell, top=85, start=115, bottom=85, end=115)
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(left.paragraphs[0].add_run("ЧАСТЬ III\n"), size=6.8, bold=True, color=WHITE, family=UI_FONT)
    _font(left.paragraphs[0].add_run("09"), size=10.0, bold=True, color=WHITE, family=UI_FONT)
    mp = middle.paragraphs[0]
    _font(mp.add_run("РЕКВИЗИТЫ, ОПЛАТА И ПОДПИСИ\n"), size=8.8, bold=True, color=NAVY, family=UI_FONT)
    _font(mp.add_run("Идентификация сторон и способы оплаты"), size=6.5, color=MUTED, family=UI_FONT)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(rp.add_run("ДОГОВОР № {{ contract_number }}\n"), size=7.0, bold=True, color=NAVY, family=UI_FONT)
    _font(rp.add_run("{{ contract_date }} · ZAKONEXPERT"), size=6.3, color=MUTED, family=UI_FONT)
    _keep(hero.rows[0])

    parties = doc.add_table(rows=1, cols=2)
    _table(parties, [CONTENT_W_MM / 2, CONTENT_W_MM / 2], border=BORDER, size=6)
    _shade(parties.cell(0, 0), PALE)
    _shade(parties.cell(0, 1), PALE)
    _card_lines(
        parties.cell(0, 0),
        "ИСПОЛНИТЕЛЬ",
        [
            ("{{ executor_brand_name }}", True),
            ("{{ executor_full_name }}", False),
            ("{{ executor_identifier_label }}: {{ executor_identifier }}", False),
            ("Руководитель: {{ executor_director_name }}", False),
            ("Юридический адрес: {{ executor_address }}", False),
            ("Тел./WhatsApp: {{ executor_phone }}", False),
            ("Сайт: {{ executor_website }}", False),
        ],
    )
    _card_lines(
        parties.cell(0, 1),
        "КЛИЕНТ",
        [
            ("{{ client_full_name }}", True),
            ("ИИН: {{ client_iin }}", False),
            ("Тел./WhatsApp: {{ client_phone }}", False),
            ("Адрес: {{ client_address }}", False),
        ],
    )
    for cell in parties.rows[0].cells:
        _cell_margins(cell, top=100, start=140, bottom=100, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_row_height(parties.rows[0], 1050)
    _keep(parties.rows[0])

    pay_heading = doc.add_paragraph()
    pay_heading.paragraph_format.space_before = Pt(4)
    pay_heading.paragraph_format.space_after = Pt(1.5)
    _font(pay_heading.add_run("ПЛАТЁЖНЫЕ РЕКВИЗИТЫ"), size=8.4, bold=True, color=NAVY, family=UI_FONT)
    _font(pay_heading.add_run("  ·  сверяйте перед оплатой"), size=7.0, color=MUTED, family=UI_FONT)
    _paragraph_rule(pay_heading, color=GOLD, size=8, space=2)

    payments = doc.add_table(rows=4, cols=2)
    _table(payments, [CONTENT_W_MM / 2, CONTENT_W_MM / 2], border=BORDER, size=5)
    pay_rows = [
        (("Получатель", "{{ executor_bank_beneficiary }}", True), ("Банк", "{{ executor_bank_name }}", False)),
        (("ИИН/БИН получателя", "{{ executor_bank_beneficiary_identifier }}", False), ("БИК / SWIFT", "{{ executor_bank_bic }}", False)),
        (("IBAN KZT", "{{ executor_bank_iban }}", True), ("Назначение", "{{ executor_bank_payment_purpose }}", False)),
        (("Kaspi", "{{ executor_kaspi_number }}", True), ("Получатель Kaspi", "{{ executor_kaspi_receiver }}", False)),
    ]
    for ridx, (row, pair) in enumerate(zip(payments.rows, pay_rows, strict=True)):
        for cidx, (cell, item) in enumerate(zip(row.cells, pair, strict=True)):
            _shade(cell, WHITE if (ridx + cidx) % 2 == 0 else PALE)
            _payment_cell(cell, item[0], item[1], emphasize=item[2])
        _set_row_height(row, 450)
        _keep(row)

    safe = doc.add_table(rows=1, cols=2)
    _table(safe, [4.5, CONTENT_W_MM - 4.5], border=GOLD, size=6)
    _shade(safe.cell(0, 0), GOLD)
    _shade(safe.cell(0, 1), PALE_GOLD)
    sp = safe.cell(0, 1).paragraphs[0]
    _font(sp.add_run("БЕЗОПАСНОСТЬ ОПЛАТЫ\n"), size=6.5, bold=True, color=GOLD, family=UI_FONT)
    _font(
        sp.add_run(
            "Перед переводом сверьте получателя, ИИН/БИН, сумму и назначение с этим разделом и сохраните чек. "
            "Получатель банковского платежа может отличаться от наименования Исполнителя только если он прямо указан здесь. "
            "ZakonExpert не запрашивает SMS-коды, пароли и доступ к банковскому приложению."
        ),
        size=7.45,
    )
    for cell in safe.rows[0].cells:
        _cell_margins(cell, top=80, start=95, bottom=80, end=95)
    _keep(safe.rows[0])

    sign_heading = doc.add_paragraph()
    sign_heading.paragraph_format.space_before = Pt(4)
    sign_heading.paragraph_format.space_after = Pt(1.5)
    _font(sign_heading.add_run("ПОДПИСИ СТОРОН"), size=8.4, bold=True, color=NAVY, family=UI_FONT)
    _font(sign_heading.add_run("  ·  экземпляр становится завершённым после подписания"), size=7.0, color=MUTED, family=UI_FONT)
    _paragraph_rule(sign_heading, color=GOLD, size=8, space=2)

    signs = doc.add_table(rows=2, cols=2)
    _table(signs, [CONTENT_W_MM / 2, CONTENT_W_MM / 2], border=BORDER, size=6)
    for cell, text in zip(signs.rows[0].cells, ("ИСПОЛНИТЕЛЬ", "КЛИЕНТ"), strict=True):
        _shade(cell, NAVY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(cell.paragraphs[0].add_run(text), size=7.2, bold=True, color=WHITE, family=UI_FONT)
        _cell_margins(cell, top=55, start=70, bottom=55, end=70)
    exec_cell, client_cell = signs.rows[1].cells
    _shade(exec_cell, WHITE)
    _shade(client_cell, WHITE)
    exec_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    client_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _cell_margins(exec_cell, top=95, start=120, bottom=80, end=120)
    _cell_margins(client_cell, top=95, start=140, bottom=80, end=140)
    ep = exec_cell.paragraphs[0]
    _font(ep.add_run("Руководитель · {{ executor_signer_short_name }}"), size=7.7, bold=True, color=NAVY, family=UI_FONT)
    mark = exec_cell.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark.paragraph_format.space_before = Pt(2)
    mark.paragraph_format.space_after = Pt(0)
    mark.add_run("{{ executor_signature_block }}")
    cp = client_cell.paragraphs[0]
    _font(cp.add_run("Подпись Клиента"), size=7.7, bold=True, color=NAVY, family=UI_FONT)
    cp.paragraph_format.space_after = Pt(8)
    raw = client_cell.add_paragraph()
    raw.paragraph_format.space_after = Pt(6)
    _font(raw.add_run("{{ client_signature }}"), size=8.0)
    line = client_cell.add_paragraph()
    line.paragraph_format.space_after = Pt(5)
    _font(line.add_run("______________________"), size=8.6, color=NAVY)
    name = client_cell.add_paragraph()
    name.paragraph_format.space_after = Pt(4)
    _font(name.add_run("{{ client_full_name }}"), size=8.0, bold=True)
    date = client_cell.add_paragraph()
    date.paragraph_format.space_after = Pt(0)
    _font(date.add_run("Дата подписания: {{ client_signature_date }}"), size=7.8, color=MUTED)
    _set_row_height(signs.rows[1], 1900)
    _keep(signs.rows[1])

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(2)
    end.paragraph_format.space_after = Pt(0)
    _font(
        end.add_run(
            "Подписывая договор, Клиент подтверждает, что ознакомился с ключевыми условиями, объёмом услуг, стоимостью, "
            "порядком оплаты и правилами прекращения договора."
        ),
        size=6.7,
        italic=True,
        color=MUTED,
    )


def build_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_defaults(doc)
    doc.core_properties.title = "Премиальный договор оказания услуг ZakonExpert"
    doc.core_properties.subject = SCHEMA_MARKER
    doc.core_properties.author = "ТОО «ZakonExpert»"
    _header_footer(doc)
    _page1(doc)
    _page2(doc)
    _page3(doc)
    with tempfile.NamedTemporaryFile(prefix="master_v3_", suffix=".docx", dir=output_path.parent, delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        doc.save(str(tmp_path))
        tmp_path.replace(output_path)
    finally:
        tmp_path.unlink(missing_ok=True)
    return output_path


def ensure_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    if output_path.exists():
        try:
            current = Document(str(output_path))
            if current.core_properties.subject == SCHEMA_MARKER:
                return output_path
        except (OSError, ValueError, KeyError):
            pass
    return build_master_template(output_path)
