from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from pypdf import PdfReader

_SPACE_RE = re.compile(r"[\t\u00a0 ]+")
_PHONE_RE = re.compile(r"\+?7[\d\s()\-]{9,24}")


def _clean(value: str | None) -> str:
    return _SPACE_RE.sub(" ", str(value or "")).strip(" \n\r\t·•|;,")


def _all_docx_text(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    chunks: list[str] = []
    chunks.extend(p.text for p in document.paragraphs if p.text)
    for table in document.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells if cell.text))
    for section in document.sections:
        for part in (section.header, section.footer):
            chunks.extend(p.text for p in part.paragraphs if p.text)
            for table in part.tables:
                for row in table.rows:
                    chunks.append(" | ".join(cell.text for cell in row.cells if cell.text))
    return "\n".join(chunks)


def _all_pdf_text(data: bytes) -> str:
    chunks: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(data))
        chunks = [(page.extract_text() or "") for page in reader.pages]
    except Exception:  # pypdf is the first, conservative extraction pass
        chunks = []
    text = "\n".join(chunks).strip()
    if len(text) >= 80:
        return text
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text("text") for page in doc)
    except Exception:
        return text


def extract_contract_text(data: bytes, *, filename: str, mime_type: str = "") -> str:
    suffix = Path(filename or "").suffix.lower()
    mime = (mime_type or "").lower()
    if suffix == ".docx" or "wordprocessingml" in mime:
        return _all_docx_text(data)
    if suffix == ".pdf" or mime == "application/pdf":
        return _all_pdf_text(data)
    if suffix in {".txt", ".text"} or mime.startswith("text/"):
        return data.decode("utf-8", errors="replace")
    raise ValueError("UNSUPPORTED_CONTRACT_FILE")


def _one(pattern: str, text: str, flags: int = re.IGNORECASE | re.DOTALL) -> str:
    match = re.search(pattern, text, flags)
    return _clean(match.group(1)) if match else ""


def _normalize_phone(value: str) -> str:
    value = _clean(value)
    if not value:
        return ""
    digits = re.sub(r"\D", "", value)
    if len(digits) == 11 and digits.startswith("8"):
        digits = "7" + digits[1:]
    if len(digits) == 10:
        digits = "7" + digits
    return "+" + digits if len(digits) >= 10 else value


def _generated_client_identity(text: str) -> tuple[str, str, str]:
    # Current ZakonExpert template renders: КЛИЕНТ / ФИО / ИИН 12digits · phone.
    match = re.search(
        r"КЛИЕНТ\s*[|\n ]+(.{3,180}?)\s*[|\n ]+ИИН\s*[:№]?\s*(\d{12})(?:\s*[·•|]\s*(\+?7[\d\s()\-]{9,24}))?",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if not match:
        return "", "", ""
    name = _clean(match.group(1))
    # Keep the regex from swallowing a table label from the preceding cell.
    name = re.sub(r"^(?:ИСПОЛНИТЕЛЬ|КЛИЕНТ)\s+", "", name, flags=re.IGNORECASE)
    return name[:255], match.group(2), _normalize_phone(match.group(3) or "")


def parse_contract_text(text: str) -> dict[str, Any]:
    raw = str(text or "").replace("\r", "\n")
    compact = _SPACE_RE.sub(" ", raw)
    name, iin, phone = _generated_client_identity(raw)

    if not iin:
        iin = _one(r"(?:ИИН|ЖСН)\s*[:№]?\s*(\d{12})", compact)
    if not name:
        name = _one(
            r"(?:КЛИЕНТ|ЗАКАЗЧИК)\s*(?:[:\-]|\n|\|)+\s*([^\n|]{5,180}?)(?=\s+(?:ИИН|ЖСН|удостовер|телефон|тел\.|адрес)\b)",
            raw,
        )
    if not name and iin:
        name = _one(
            rf"([А-ЯЁӘҒҚҢӨҰҮҺ][А-Яа-яЁёӘәҒғҚқҢңӨөҰұҮүҺһ\-]+(?:\s+[А-ЯЁӘҒҚҢӨҰҮҺ][А-Яа-яЁёӘәҒғҚқҢңӨөҰұҮүҺһ\-]+){{1,3}})\s*,?\s*(?:ИИН|ЖСН)\s*[:№]?\s*{re.escape(iin)}",
            compact,
        )

    if not phone:
        phone = _normalize_phone(
            _one(r"(?:Тел(?:ефон)?\.?|WhatsApp|Моб\.?)[\s/:\-]*((?:\+?7|8)[\d\s()\-]{9,24})", compact)
        )

    number = _one(
        r"(?:НОМЕР\s+ДОГОВОРА\s*)?№\s*([A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9./\-]{0,40})",
        compact,
    )
    date = _one(r"(?:ДАТА\s*)?(\d{2}[./]\d{2}[./]\d{4})(?:\s*г\.?)?", compact)
    date = date.replace(".", "-") if date else ""
    if date and re.fullmatch(r"\d{2}-\d{2}-\d{4}", date):
        day, month, year = date.split("-")
        date = f"{year}-{month}-{day}"

    service = _one(
        r"(?:УСЛУГА|ПРЕДМЕТ\s+ДОГОВОРА)\s*[|:\-]*\s*(.{5,1200}?)(?=\s*(?:\||\n)?\s*(?:02\s*)?(?:РЕЗУЛЬТАТ|СРОК|СТОИМОСТЬ|2\.))",
        raw,
    )
    if service:
        service = re.sub(r"\s+", " ", service)[:1200]

    amount_text = _one(
        r"(?:СТОИМОСТЬ|ЦЕНА\s+(?:ДОГОВОРА|УСЛУГ)?|СТОИМОСТЬ\s+УСЛУГ)[\s|:\-]*([0-9][0-9\s\u00a0]{0,18})(?=\s*(?:тенге|тг|₸|KZT))",
        compact,
    )
    amount = int(re.sub(r"\D", "", amount_text) or 0)

    address = _one(
        r"Адрес(?:\s+Клиента)?\s*[:\-]\s*(.{5,320}?)(?=\s*(?:\||\n|Тел(?:ефон)?\.?|WhatsApp|ИИН|ЖСН|ПОДПИС|$))",
        raw,
    )
    document_number = _one(
        r"(?:Удостоверение\s+личности|№\s*удостоверения|Документ)\s*(?:№|:)?\s*([A-Za-zА-Яа-я0-9\-]{5,32})",
        compact,
    )

    payment_type = ""
    lowered = compact.lower()
    if "после достижения результата" in lowered or "после результата" in lowered:
        payment_type = "after_result"
    elif "до начала оказания услуг" in lowered or "до начала работ" in lowered or "предоплат" in lowered:
        payment_type = "prepayment"
    elif "два этапа" in lowered or "50/50" in lowered:
        payment_type = "split"
    elif "полностью оплатил" in lowered or "уже оплач" in lowered:
        payment_type = "already_paid"

    return {
        "name": name[:255],
        "iin": iin[:12],
        "phone": phone[:32],
        "address": address[:512],
        "documentNumber": document_number[:64],
        "number": number[:80],
        "date": date,
        "amount": amount,
        "currency": "KZT",
        "service": service,
        "paymentType": payment_type,
        "textLength": len(raw),
    }


def parse_contract_bytes(data: bytes, *, filename: str, mime_type: str = "") -> dict[str, Any]:
    if not data:
        raise ValueError("EMPTY_CONTRACT_FILE")
    text = extract_contract_text(data, filename=filename, mime_type=mime_type)
    if len(text.strip()) < 20:
        raise ValueError("CONTRACT_TEXT_NOT_FOUND")
    result = parse_contract_text(text)
    result["filename"] = Path(filename or "contract").name[:240]
    return result
