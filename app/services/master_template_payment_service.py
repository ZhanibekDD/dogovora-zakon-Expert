from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.services.master_template_light_service import build_master_template as build_v5_template

SCHEMA_MARKER = "ZakonExpert contract schema v6 flexible-payment"
NAVY = "20364F"
GOLD = "B78A43"
PALE_GOLD = "FBF8F1"
MAIN_PHONE = "+7 700 309 7566"
KASPI_PHONE = "+7 705 876 27 95"
KASPI_RECIPIENT = "Жанибек К."


def _set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = NAVY) -> None:
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for part in (section.header, section.footer):
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _replace_text(doc, old: str, new: str) -> None:
    for paragraph in _all_paragraphs(doc):
        if old not in paragraph.text:
            continue
        for run in paragraph.runs:
            run.text = run.text.replace(old, new)


def _rewrite_payment_choices(doc) -> None:
    """Show two simple payment options without presenting Kaspi as the company phone."""
    target = None
    for table in doc.tables:
        text = _table_text(table)
        if "ОСНОВНОЙ СПОСОБ" in text and "ПОДТВЕРЖДЕНИЕ" in text:
            target = table
            break
    if target is None:
        return

    rows = [
        (
            "ПО СЧЁТУ",
            "Оплата по счёту или платёжной ссылке ZakonExpert для конкретного договора.",
        ),
        (
            "ЧЕРЕЗ KASPI",
            f"Если Клиенту удобнее, по согласованию с Исполнителем можно оплатить переводом через Kaspi по номеру {KASPI_PHONE}. Получатель: {KASPI_RECIPIENT}",
        ),
        (
            "ПОДТВЕРЖДЕНИЕ",
            "После оплаты сохраните чек или иной платёжный документ, подтверждающий сумму и дату перевода.",
        ),
        (
            "КОНТАКТ",
            f"Для получения счёта и по вопросам оплаты: {MAIN_PHONE} · zakonexpertt.kz.",
        ),
    ]
    for row, (label, value) in zip(target.rows, rows, strict=False):
        cells = row.cells
        _set_cell_fill(cells[0], PALE_GOLD)
        _set_cell_text(cells[0], label, bold=True, color=GOLD)
        # In v5 columns 1..3 are already merged into one value cell.
        _set_cell_text(cells[1], value, color=NAVY)


def _remove_security_notice(doc) -> None:
    for table in list(doc.tables):
        if "БЕЗОПАСНОСТЬ ОПЛАТЫ" not in _table_text(table):
            continue
        element = table._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def build_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    build_v5_template(output_path)
    doc = Document(str(output_path))
    doc.core_properties.subject = SCHEMA_MARKER
    doc.core_properties.title = "Договор оказания услуг ZakonExpert — flexible payment"

    _replace_text(
        doc,
        "по счёту или платёжной ссылке, письменно подтверждённой Исполнителем",
        f"по счёту или платёжной ссылке ZakonExpert, а по согласованию с Исполнителем — переводом через Kaspi по номеру {KASPI_PHONE}",
    )
    _rewrite_payment_choices(doc)
    _remove_security_notice(doc)

    doc.save(str(output_path))
    return output_path


def ensure_master_template(output_path: Path) -> Path:
    output_path = Path(output_path)
    if output_path.exists():
        try:
            current = Document(str(output_path))
            if current.core_properties.subject == SCHEMA_MARKER:
                return output_path
        except (OSError, ValueError, KeyError):
            pass
    return build_master_template(output_path)
