from __future__ import annotations

from docx import Document as DocxReader
from docx.shared import Mm

from app.core.config import get_settings
from app.services.master_template_payment_service import SCHEMA_MARKER, ensure_master_template


def _load_master_template():
    settings = get_settings()
    path = ensure_master_template(settings.templates_dir / "master_v1.docx")
    return DocxReader(str(path))


def _table_text(table) -> str:
    chunks: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            chunks.append(cell.text)
            chunks.extend(_table_text(nested) for nested in cell.tables)
    return "\n".join(chunks)


def _all_text(doc) -> str:
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(_table_text(table) for table in doc.tables)
    return text


def test_page_is_a4_with_adequate_margins() -> None:
    doc = _load_master_template()
    section = doc.sections[0]
    assert abs(section.page_width - Mm(210)) < Mm(1)
    assert abs(section.page_height - Mm(297)) < Mm(1)
    tolerance = Mm(0.2)
    assert section.top_margin + tolerance >= Mm(14)
    assert section.bottom_margin + tolerance >= Mm(14)
    assert section.left_margin + tolerance >= Mm(14)
    assert section.right_margin + tolerance >= Mm(14)


def test_body_font_is_times_new_roman_11_or_12pt() -> None:
    doc = _load_master_template()
    normal_style = doc.styles["Normal"]
    assert normal_style.font.name == "Times New Roman"
    assert normal_style.font.size.pt in (11, 12)


def test_official_brand_schema_and_branding_are_present() -> None:
    doc = _load_master_template()
    all_text = _all_text(doc)
    assert doc.core_properties.subject == SCHEMA_MARKER
    assert "v6 flexible-payment" in SCHEMA_MARKER
    assert "ИНДИВИДУАЛЬНЫЙ ДОГОВОР" in all_text
    assert "КЛЮЧЕВЫЕ УСЛОВИЯ" in all_text
    assert "+7 700 309 7566" in all_text
    assert "{{ executor_phone }}" not in all_text

    header = doc.sections[0].header
    image_parts = [part for part in header.part.related_parts.values() if part.content_type.startswith("image/")]
    assert image_parts


def test_signature_area_table_present_with_two_columns() -> None:
    doc = _load_master_template()
    assert len(doc.tables) >= 4
    signature_table = doc.tables[-1]
    assert len(signature_table.columns) == 2


def test_signature_placeholders_present_in_template() -> None:
    doc = _load_master_template()
    full_text = "\n".join(_table_text(table) for table in doc.tables)
    assert "{{ executor_signature_block }}" in full_text
    assert "{{ client_signature }}" in full_text
    assert "{{ client_signature_date }}" in full_text


def test_three_part_contract_and_flexible_payment_flow_are_present() -> None:
    doc = _load_master_template()
    all_text = _all_text(doc)
    assert "ЧАСТЬ I" in all_text
    assert "ЧАСТЬ II" in all_text
    assert "ЧАСТЬ III" in all_text
    assert "ПОРЯДОК ОПЛАТЫ" in all_text
    assert "Оплата по счёту или платёжной ссылке" in all_text
    assert "ЧЕРЕЗ KASPI" in all_text
    assert "+7 705 876 27 95" in all_text
    assert "Жанибек К." in all_text
    assert "БЕЗОПАСНОСТЬ ОПЛАТЫ" not in all_text
    assert "SMS-коды" not in all_text
    assert "{{ executor_bank_beneficiary }}" not in all_text
    assert "{{ executor_bank_beneficiary_identifier }}" not in all_text
    assert "{{ executor_bank_iban }}" not in all_text
    assert "{{ executor_bank_bic }}" not in all_text


def test_required_placeholders_present() -> None:
    doc = _load_master_template()
    all_text = _all_text(doc)
    required = [
        "{{ contract_number }}",
        "{{ contract_date }}",
        "{{ contract_city }}",
        "{{ client_full_name }}",
        "{{ client_iin }}",
        "{{ client_phone }}",
        "{{ service_subject }}",
        "{{ result_definition }}",
        "{{ amount_digits }}",
        "{{ amount_words }}",
        "{{ payment_terms }}",
        "{{ penalty_clause }}",
        "{{ executor_website }}",
    ]
    for placeholder in required:
        assert placeholder in all_text, f"missing {placeholder}"
