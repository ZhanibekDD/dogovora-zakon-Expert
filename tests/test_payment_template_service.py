from __future__ import annotations

from docx import Document

from app.services.master_template_payment_service import (
    KASPI_PHONE,
    KASPI_RECIPIENT,
    MAIN_PHONE,
    _remove_security_notice,
    _rewrite_payment_choices,
)


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_payment_choice_rewrite_keeps_main_phone_separate_from_kaspi() -> None:
    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    labels = ["ОСНОВНОЙ СПОСОБ", "ИДЕНТИФИКАЦИЯ", "ПОДТВЕРЖДЕНИЕ", "КОНТАКТ"]
    for row, label in zip(table.rows, labels, strict=True):
        row.cells[0].text = label
        row.cells[1].merge(row.cells[3])
        row.cells[1].text = "old"

    _rewrite_payment_choices(doc)
    text = _table_text(table)

    assert "ПО СЧЁТУ" in text
    assert "ЧЕРЕЗ KASPI" in text
    assert KASPI_PHONE in text
    assert KASPI_RECIPIENT in text
    assert MAIN_PHONE in text


def test_security_payment_notice_is_removed_entirely() -> None:
    doc = Document()
    notice = doc.add_table(rows=1, cols=1)
    notice.cell(0, 0).text = "БЕЗОПАСНОСТЬ ОПЛАТЫ: старый текст"
    signature = doc.add_table(rows=1, cols=2)
    signature.cell(0, 0).text = "ИСПОЛНИТЕЛЬ"
    signature.cell(0, 1).text = "КЛИЕНТ"

    _remove_security_notice(doc)
    text = "\n".join(_table_text(table) for table in doc.tables)

    assert "БЕЗОПАСНОСТЬ ОПЛАТЫ" not in text
    assert "ИСПОЛНИТЕЛЬ" in text
    assert "КЛИЕНТ" in text
