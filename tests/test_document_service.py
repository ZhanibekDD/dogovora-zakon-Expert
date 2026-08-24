from __future__ import annotations

from pathlib import Path

from docx import Document as DocxReader

from app.core.config import get_settings
from app.schemas.contract import ContractRenderContext
from app.services import document_service


def _sample_context(**overrides) -> ContractRenderContext:
    base = dict(
        contract_number=42,
        contract_date="13.07.2026",
        contract_city="г. Талдыкорган",
        client_full_name="СЕЙТЖАНОВ АЙБЕК НҰРЛАНҰЛЫ",
        client_iin="010312500019",
        client_phone="+7 701 234 5678",
        client_address="не указан",
        service_subject="Тестовый предмет договора.",
        service_actions="тестовые действия",
        result_definition="тестовый результат",
        amount_digits="120 000 тенге",
        amount_words="сто двадцать тысяч тенге",
        payment_terms="Оплата производится в день подписания договора.",
        work_period="до 30 календарных дней",
        penalty_clause="Пеня 0,1% в день.",
        executor_name="ТОО «ZakonExpert»",
        executor_full_name="Товарищество с ограниченной ответственностью «ZakonExpert»",
        executor_brand_name="ТОО «ZakonExpert»",
        executor_identifier_label="БИН",
        executor_identifier="260740044168",
        executor_director_name="Кияшев Жанибек Даулетович",
        executor_signer_short_name="Кияшев Ж.Д.",
        executor_phone="+7 700 309 7566",
        executor_address="г. Талдыкорган, ул. Акын Сара, 152",
        executor_website="zakonexpertt.kz",
        executor_payment_details="по реквизитам раздела 9",
        executor_bank_beneficiary="Жанибек Кияшев Даулетович",
        executor_bank_beneficiary_identifier="000725500183",
        executor_bank_name="АО «Фридом Банк Казахстан»",
        executor_bank_bic="KSNVKZKA",
        executor_bank_iban="KZ95551V600001202152",
        executor_bank_payment_purpose="Оплата по договору № 42",
        executor_kaspi_number="+7 705 876 27 95",
        executor_kaspi_receiver="Жанибек К.",
    )
    base.update(overrides)
    return ContractRenderContext(**base)


def test_render_draft_docx_contains_client_data(tmp_path: Path) -> None:
    settings = get_settings()
    template_path = settings.templates_dir / "master_v1.docx"
    output_path = tmp_path / "draft.docx"

    document_service.render_contract_docx(
        template_docx_path=template_path,
        context=_sample_context(),
        output_path=output_path,
        include_executor_signature=False,
    )

    assert output_path.exists()
    doc = DocxReader(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs) + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "СЕЙТЖАНОВ АЙБЕК НҰРЛАНҰЛЫ" in full_text
    assert "010312500019" in full_text
    assert "120 000 тенге" in full_text
    # Raw bank beneficiary/IBAN is not printed as company requisites. The contract offers
    # invoice/payment-link payment and, if more convenient, a separate Kaspi payment option.
    assert "KZ95551V600001202152" not in full_text
    assert "zakonexpertt.kz" in full_text
    assert "+7 700 309 7566" in full_text
    assert "+7 705 876 27 95" in full_text
    assert "Жанибек К." in full_text
    assert "БЕЗОПАСНОСТЬ ОПЛАТЫ" not in full_text
    assert "{{" not in full_text


def test_client_signature_is_always_blank_regardless_of_input(tmp_path: Path) -> None:
    output_path = tmp_path / "draft_with_attempted_signature.docx"
    context = _sample_context(
        client_signature="SHOULD_NEVER_APPEAR",
        client_signature_date="01.01.2026",
    )

    settings = get_settings()
    document_service.render_contract_docx(
        template_docx_path=settings.templates_dir / "master_v1.docx",
        context=context,
        output_path=output_path,
        include_executor_signature=False,
    )

    doc = DocxReader(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs) + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "SHOULD_NEVER_APPEAR" not in full_text


def test_executor_signature_block_embedded_only_when_requested(tmp_path: Path) -> None:
    settings = get_settings()
    with_sig_path = tmp_path / "final.docx"
    without_sig_path = tmp_path / "draft.docx"

    document_service.render_contract_docx(
        template_docx_path=settings.templates_dir / "master_v1.docx",
        context=_sample_context(),
        output_path=with_sig_path,
        include_executor_signature=True,
    )
    document_service.render_contract_docx(
        template_docx_path=settings.templates_dir / "master_v1.docx",
        context=_sample_context(),
        output_path=without_sig_path,
        include_executor_signature=False,
    )

    with_sig_images = len(DocxReader(str(with_sig_path)).inline_shapes)
    without_sig_images = len(DocxReader(str(without_sig_path)).inline_shapes)
    assert with_sig_images == 1
    assert without_sig_images == 0

    widths_mm = sorted(shape.width / 36000 for shape in DocxReader(str(with_sig_path)).inline_shapes)
    expected_width = settings.executor_signature_block_width_mm
    assert any(abs(width - expected_width) < 0.5 for width in widths_mm)
