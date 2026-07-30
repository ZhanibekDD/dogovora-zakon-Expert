from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "app" / "templates" / "contracts" / "master_v1.docx"

# Design system: standard_business_brief with named legal-contract overrides.
# A4 / 18 mm margins / Times New Roman / restrained blue hierarchy are encoded here
# explicitly so Word and LibreOffice render the same geometry.
FONT = "Times New Roman"
NAVY = "000000"
MUTED = "667085"
LIGHT_BLUE = "EAF0F7"
LIGHT_BORDER = "B8C4D1"
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_MM = 18
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_MM * 2
CONTENT_WIDTH_TWIPS = round(CONTENT_WIDTH_MM / 25.4 * 1440)
HALF_WIDTHS_TWIPS = (CONTENT_WIDTH_TWIPS // 2, CONTENT_WIDTH_TWIPS - CONTENT_WIDTH_TWIPS // 2)


def _set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = f"w:{edge}"
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_twips: tuple[int, ...]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_twips[min(index, len(widths_twips) - 1)]
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.width = Twips(width)


def _set_table_borders(table, *, color: str = LIGHT_BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _keep_row_together(row) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True


def _body(document: DocumentType, text: str, *, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Mm(8)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, bold=True)
        text = text[len(bold_lead) :]
    run = paragraph.add_run(text)
    _set_run_font(run)
    return paragraph


def _section_title(document: DocumentType, number: str, title: str):
    paragraph = document.add_paragraph(style="ZE Section")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number}. {title}")
    _set_run_font(run, size=10.5, bold=True, color=NAVY)
    return paragraph


def _plain_cell_text(cell, lines: list[tuple[str, bool]]) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, (text, bold) in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        _set_run_font(paragraph.add_run(text), size=9.5, bold=bold)


def _build_header_footer(document: DocumentType) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].clear()
    table = header.add_table(rows=1, cols=2, width=Mm(CONTENT_WIDTH_MM))
    _set_table_geometry(table, HALF_WIDTHS_TWIPS)
    _remove_table_borders(table)
    left, right = table.rows[0].cells
    _set_run_font(left.paragraphs[0].add_run("ZAKONEXPERT"), size=8.5, bold=True, color=NAVY)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        right_p.add_run("Договор оказания услуг № {{ contract_number }}"),
        size=8.5,
        color=MUTED,
    )

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    _set_run_font(
        footer_p.add_run("ТОО «ZakonExpert» | {{ executor_phone }} | zakonexpertt.kz"),
        size=8,
        color=MUTED,
    )


def _build_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    section_style = document.styles.add_style("ZE Section", WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = normal
    section_style.font.name = FONT
    section_style.font.size = Pt(10.5)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor.from_string(NAVY)
    section_style.paragraph_format.space_before = Pt(9)
    section_style.paragraph_format.space_after = Pt(4)
    section_style.paragraph_format.keep_with_next = True


def _build_parties_page(document: DocumentType) -> None:
    document.add_page_break()
    _section_title(document, "9", "РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")

    requisites = document.add_table(rows=2, cols=2)
    _set_table_geometry(requisites, HALF_WIDTHS_TWIPS)
    _set_table_borders(requisites)
    headers = requisites.rows[0].cells
    for cell, text in zip(headers, ("ИСПОЛНИТЕЛЬ", "КЛИЕНТ"), strict=True):
        _shade_cell(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(text), size=10, bold=True, color=NAVY)
    _set_repeat_table_header(requisites.rows[0])

    details = requisites.rows[1].cells
    _plain_cell_text(
        details[0],
        [
            ("{{ executor_brand_name }}", True),
            ("{{ executor_full_name }}", False),
            ("{{ executor_identifier_label }}: {{ executor_identifier }}", False),
            ("Руководитель: {{ executor_director_name }}", False),
            ("Адрес: {{ executor_address }}", False),
            ("Тел./WhatsApp: {{ executor_phone }}", False),
        ],
    )
    _plain_cell_text(
        details[1],
        [
            ("{{ client_full_name }}", True),
            ("ИИН: {{ client_iin }}", False),
            ("Телефон/WhatsApp: {{ client_phone }}", False),
            ("Адрес: {{ client_address }}", False),
        ],
    )
    for cell in details:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _keep_row_together(requisites.rows[1])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

    signatures = document.add_table(rows=2, cols=2)
    _set_table_geometry(signatures, HALF_WIDTHS_TWIPS)
    _set_table_borders(signatures)
    sig_headers = signatures.rows[0].cells
    for cell, text in zip(sig_headers, ("ПОДПИСЬ И ПЕЧАТЬ ИСПОЛНИТЕЛЯ", "ПОДПИСЬ КЛИЕНТА"), strict=True):
        _shade_cell(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(text), size=9.5, bold=True, color=NAVY)

    executor_cell, client_cell = signatures.rows[1].cells
    executor_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    client_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_margins(executor_cell, top=140, start=140, bottom=140, end=140)
    _set_cell_margins(client_cell, top=140, start=140, bottom=140, end=140)

    exec_label = executor_cell.paragraphs[0]
    exec_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    exec_label.paragraph_format.space_after = Pt(3)
    _set_run_font(exec_label.add_run("Руководитель"), size=9.5, bold=True)

    executor_mark = executor_cell.add_paragraph()
    executor_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    executor_mark.paragraph_format.space_before = Pt(0)
    executor_mark.paragraph_format.space_after = Pt(0)
    executor_mark.paragraph_format.line_spacing = 1
    executor_mark.add_run("{{ executor_signature_block }}")

    client_label = client_cell.paragraphs[0]
    client_label.paragraph_format.space_after = Pt(4)
    _set_run_font(client_label.add_run("Подпись Клиента:"), size=9.5, bold=True)

    client_image = client_cell.add_paragraph()
    client_image.paragraph_format.space_before = Pt(0)
    client_image.paragraph_format.space_after = Pt(0)
    client_image.paragraph_format.line_spacing = 1
    _set_run_font(client_image.add_run("{{ client_signature }}"), size=9.5)
    _set_run_font(client_image.add_run("\n\n"), size=9.5)

    client_line = client_cell.add_paragraph()
    client_line.paragraph_format.space_before = Pt(0)
    client_line.paragraph_format.space_after = Pt(5)
    _set_run_font(client_line.add_run("________________ / {{ client_full_name }} /"), size=9.5)

    client_date = client_cell.add_paragraph()
    client_date.paragraph_format.space_before = Pt(0)
    client_date.paragraph_format.space_after = Pt(0)
    _set_run_font(
        client_date.add_run("Дата подписания: {{ client_signature_date }}"),
        size=9.5,
    )
    _keep_row_together(signatures.rows[1])


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(9)

    document.core_properties.title = "Договор оказания услуг ZakonExpert"
    document.core_properties.subject = "Мастер-шаблон договора"
    document.core_properties.author = "ТОО «ZakonExpert»"
    _build_styles(document)
    _build_header_footer(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(2)
    _set_run_font(
        title.add_run("ДОГОВОР ОКАЗАНИЯ УСЛУГ"),
        size=15,
        bold=True,
        color=NAVY,
    )

    number = document.add_paragraph()
    number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number.paragraph_format.space_before = Pt(0)
    number.paragraph_format.space_after = Pt(7)
    _set_run_font(number.add_run("№ {{ contract_number }}"), size=11.5, bold=True)

    metadata = document.add_table(rows=1, cols=2)
    _set_table_geometry(metadata, HALF_WIDTHS_TWIPS)
    _remove_table_borders(metadata)
    city_p = metadata.cell(0, 0).paragraphs[0]
    _set_run_font(city_p.add_run("{{ contract_city }}"), size=10.5, bold=True)
    date_p = metadata.cell(0, 1).paragraphs[0]
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(date_p.add_run("{{ contract_date }} г."), size=10.5, bold=True)

    _body(
        document,
        "{{ executor_full_name }}, {{ executor_identifier_label }} {{ executor_identifier }}, "
        "в лице руководителя {{ executor_director_name }}, действующего на основании Устава, "
        "именуемое в дальнейшем «Исполнитель», с одной стороны, и гражданин(ка) "
        "{{ client_full_name }}, ИИН {{ client_iin }}, именуемый(ая) в дальнейшем «Клиент», "
        "с другой стороны, совместно именуемые «Стороны», заключили настоящий договор "
        "о нижеследующем.",
    )

    _section_title(document, "1", "ПРЕДМЕТ ДОГОВОРА")
    _body(document, "1.1. {{ service_subject }}", bold_lead="1.1. ")
    _body(
        document,
        "1.2. В состав услуг входят: {{ service_actions }}.",
        bold_lead="1.2. ",
    )
    _body(
        document,
        "1.3. Ожидаемый и проверяемый результат оказания услуг: {{ result_definition }}.",
        bold_lead="1.3. ",
    )

    _section_title(document, "2", "ПОРЯДОК И СРОКИ ОКАЗАНИЯ УСЛУГ")
    _body(
        document,
        "2.1. Исполнитель приступает к оказанию услуг после получения от Клиента документов "
        "и сведений, необходимых для выполнения поручения, а при предусмотренной предоплате - "
        "также после её внесения.",
        bold_lead="2.1. ",
    )
    _body(
        document,
        "2.2. Срок оказания услуг: {{ work_period }}. Сроки рассмотрения обращений судом, "
        "нотариусом, ЧСИ, банком, взыскателем и государственными органами не зависят от "
        "Исполнителя и в срок оказания услуг не включаются.",
        bold_lead="2.2. ",
    )
    _body(
        document,
        "2.3. Клиент обязан незамедлительно передавать Исполнителю поступившие уведомления, "
        "постановления, судебные извещения, ответы государственных органов и организаций и "
        "иные документы, имеющие отношение к делу.",
        bold_lead="2.3. ",
    )

    _section_title(document, "3", "СТОИМОСТЬ УСЛУГ И ПОРЯДОК ОПЛАТЫ")
    _body(
        document,
        "3.1. Общая стоимость услуг по настоящему договору составляет {{ amount_digits }} "
        "({{ amount_words }}).",
        bold_lead="3.1. ",
    )
    _body(document, "3.2. {{ payment_terms }}", bold_lead="3.2. ")
    _body(
        document,
        "3.3. Оплата производится {{ executor_payment_details }}. Факт оплаты подтверждается "
        "банковским чеком, квитанцией, распиской либо иным платёжным документом.",
        bold_lead="3.3. ",
    )
    _body(document, "3.4. {{ penalty_clause }}", bold_lead="3.4. ")

    document.add_page_break()
    _section_title(document, "4", "ПРАВА И ОБЯЗАННОСТИ СТОРОН")
    _body(
        document,
        "4.1. Исполнитель обязуется добросовестно подготовить документы, информировать "
        "Клиента о существенных действиях и передавать Клиенту полученные ответы и документы.",
        bold_lead="4.1. ",
    )
    _body(
        document,
        "4.2. Исполнитель вправе выбирать законный способ защиты интересов Клиента, "
        "запрашивать дополнительные документы и приостанавливать работу до их предоставления.",
        bold_lead="4.2. ",
    )
    _body(
        document,
        "4.3. Клиент обязуется предоставлять достоверные и полные сведения, своевременно "
        "подписывать подготовленные документы, оплачивать обязательные государственные и "
        "нотариальные расходы, лично являться в суд, к нотариусу или ЧСИ, если это требуется.",
        bold_lead="4.3. ",
    )
    _body(
        document,
        "4.4. Клиент подтверждает, что сообщил Исполнителю обо всех известных исполнительных "
        "производствах, судебных актах, соглашениях, платежах, ранее поданных заявлениях и "
        "иных обстоятельствах, влияющих на оказание услуг.",
        bold_lead="4.4. ",
    )

    _section_title(document, "5", "ПРИЁМКА УСЛУГ")
    _body(
        document,
        "5.1. После достижения предусмотренного договором результата Исполнитель направляет "
        "Клиенту итоговый документ или уведомление о результате посредством Telegram, "
        "WhatsApp или электронной почты.",
        bold_lead="5.1. ",
    )
    _body(
        document,
        "5.2. Если Клиент в течение 3 (трёх) календарных дней не направит мотивированные "
        "письменные возражения с указанием конкретных недостатков, услуги считаются оказанными "
        "и принятыми в полном объёме.",
        bold_lead="5.2. ",
    )

    _section_title(document, "6", "ОТВЕТСТВЕННОСТЬ И РАЗРЕШЕНИЕ СПОРОВ")
    _body(
        document,
        "6.1. Стороны несут ответственность в соответствии с законодательством Республики "
        "Казахстан и настоящим договором.",
        bold_lead="6.1. ",
    )
    _body(
        document,
        "6.2. Исполнитель не отвечает за незаконные либо несвоевременные действия или "
        "бездействие нотариуса, ЧСИ, суда, взыскателя, банка и государственных органов, но "
        "обязуется подготовить и передать Клиенту предусмотренные договором документы.",
        bold_lead="6.2. ",
    )
    _body(
        document,
        "6.3. Стороны признают юридическую силу переписки и документов, переданных через "
        "WhatsApp, Telegram, SMS, электронную почту и иные согласованные средства связи, "
        "если возможно достоверно определить отправителя, содержание и дату сообщения.",
        bold_lead="6.3. ",
    )
    _body(
        document,
        "6.4. До обращения в суд заинтересованная Сторона направляет письменную претензию. "
        "Срок ответа - 5 (пять) рабочих дней с момента её получения. При недостижении "
        "соглашения спор рассматривается в установленном законом порядке.",
        bold_lead="6.4. ",
    )

    _section_title(document, "7", "ПЕРСОНАЛЬНЫЕ ДАННЫЕ И УВЕДОМЛЕНИЯ")
    _body(
        document,
        "7.1. Подписывая настоящий договор, Клиент даёт согласие на сбор и обработку "
        "Исполнителем своих персональных данных в объёме, необходимом для исполнения "
        "настоящего договора.",
        bold_lead="7.1. ",
    )
    _body(
        document,
        "7.2. Уведомления направляются по телефону, WhatsApp, Telegram или электронной почте, "
        "указанным Сторонами, и считаются полученными в день отправки.",
        bold_lead="7.2. ",
    )

    _section_title(document, "8", "ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
    _body(
        document,
        "8.1. Настоящий договор вступает в силу с даты его подписания и действует до полного "
        "исполнения Сторонами обязательств.",
        bold_lead="8.1. ",
    )
    _body(
        document,
        "8.2. Изменения и дополнения действительны при письменном согласовании Сторонами, "
        "включая согласование через WhatsApp, Telegram или электронную переписку.",
        bold_lead="8.2. ",
    )
    _body(
        document,
        "8.3. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "либо подписан посредством простой электронной подписи Клиента.",
        bold_lead="8.3. ",
    )

    _build_parties_page(document)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Master template written to {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
    sys.exit(0)
