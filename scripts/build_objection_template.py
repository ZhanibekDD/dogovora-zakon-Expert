"""Builds the docxtpl template for 'Возражение на исполнительную надпись нотариуса', based on
the real, already-in-use objection documents of IP ZakonExpert. All genuinely variable text
(gendered clauses, optional email line, optional obligation-basis paragraph) is precomputed in
Python by app/services/objection_service.py and passed in as plain strings - the template
itself only holds the fixed structure and placeholders, mirroring how the main contract
template (scripts/build_master_template.py) is built.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

OUTPUT_PATH = (
    Path(__file__).resolve().parent.parent / "app" / "templates" / "objections" / "objection_v1.docx"
)

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)


def _set_base_style(document: DocumentObject) -> None:
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)

    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)


def _add_page_numbers(document: DocumentObject) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._r.append(fld_begin)
    instr = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = "PAGE"
    run._r.append(instr)
    fld_end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(fld_end)


def _p(document: DocumentObject, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold


def build() -> None:
    document = Document()
    _set_base_style(document)
    _add_page_numbers(document)

    _p(document, "Нотариусу г. {{ notary_city }}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "{{ notary_full_name }}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "Лицензия №{{ notary_license_number }} от {{ notary_license_date }}", align=WD_ALIGN_PARAGRAPH.LEFT)

    _p(document, "от: {{ client_full_name }}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "{{ client_birth_date }} г.р., ИИН {{ client_iin }}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "адрес: {{ client_address }}", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "тел.: {{ client_phone }}{{ client_email_line }}", align=WD_ALIGN_PARAGRAPH.LEFT)

    _p(document, "ВОЗРАЖЕНИЕ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _p(document, "на исполнительную надпись", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    _p(
        document,
        "{{ writ_date_long }} Вами, нотариусом г. {{ notary_city }} {{ notary_full_name }}, "
        "была совершена исполнительная надпись, уникальный номер {{ unique_number }}, "
        "зарегистрированная в реестре за №{{ registry_number }}, о взыскании с "
        "{{ client_full_name }}, {{ client_birth_date }} г.р., ИИН {{ client_iin }}, в пользу "
        "{{ creditor_name_genitive }}, БИН {{ creditor_bin }}, задолженности в сумме "
        "{{ debt_amount_text }}, расходов по совершению исполнительной надписи в сумме "
        "{{ fee_amount_text }}, всего {{ total_amount_text }}.",
    )

    _p(document, "{{ disagree_clause }}")

    _p(
        document,
        "Согласно статье 92-2 Закона Республики Казахстан «О нотариате», исполнительная "
        "надпись совершается при условии, что представленные документы подтверждают "
        "бесспорность задолженности или иной ответственности должника перед взыскателем. "
        "При наличии спора о праве, несогласия должника с требованием либо с размером "
        "задолженности требование не может считаться бесспорным.",
    )

    _p(document, "{{ not_signed_clause }}")

    _p(document, "{{ obligation_paragraph }}")

    _p(document, "{{ learned_clause }}")

    _p(
        document,
        "В соответствии со статьей 92-8 Закона Республики Казахстан «О нотариате» нотариус "
        "выносит постановление об отмене исполнительной надписи не позднее трех рабочих дней "
        "со дня поступления возражения должника. В случае если исполнительная надпись по "
        "возражению должника не отменена, ее оспаривание осуществляется в судебном порядке.",
    )

    _p(document, "На основании вышеизложенного, ПРОШУ:")
    _p(document, "1. Принять настоящее возражение против заявленного требования.")
    _p(
        document,
        "2. Отменить исполнительную надпись от {{ writ_date_long }}, уникальный номер "
        "{{ unique_number }}, зарегистрированную в реестре за №{{ registry_number }}, "
        "совершенную нотариусом г. {{ notary_city }} {{ notary_full_name }}.",
    )
    _p(
        document,
        "3. Направить постановление об отмене исполнительной надписи должнику по номеру "
        "WhatsApp: {{ client_phone }} либо иным доступным способом.",
    )
    _p(document, "4. Уведомить взыскателя и судебного исполнителя об отмене исполнительной надписи.")

    _p(document, "Приложения:", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "1. Копия исполнительной надписи.", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "2. Копия удостоверения личности должника.", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "3. Иные документы при наличии.", align=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_paragraph()
    _p(document, "{{ client_signature_line }} __________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    _p(document, "Дата: {{ objection_date }}", align=WD_ALIGN_PARAGRAPH.LEFT)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Objection template written to {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
    sys.exit(0)
